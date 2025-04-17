import os
import json
import re
import argparse
from typing import List, Optional

import torch
from transformers import pipeline, AutoTokenizer, AutoModelForCausalLM
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT



try:
    from docx2pdf import convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False


class TableFiller:
    def __init__(self, model_name: str = "HuggingFaceH4/zephyr-7b-beta"):
        print(f"Loading model: {model_name}")
        self.tokenizer = AutoTokenizer.from_pretrained(model_name)
        self.model = AutoModelForCausalLM.from_pretrained(
            model_name,
            device_map="auto",
            torch_dtype=torch.float16 if torch.cuda.is_available() else torch.float32
        )
        self.generator = pipeline(
            "text-generation",
            model=self.model,
            tokenizer=self.tokenizer,
            torch_dtype=torch.float16 if torch.cuda.is_available() else torch.float32
        )
        self.system_prompt = """
                            You are a helpful assistant that generates tables based on the provided prompt.
                            You will receive a topic, row headers, and column headers.
                            Your task is to fill in a table with realistic and plausible data.

                            Return ONLY a valid JSON array of arrays like: [[value1, value2, ...], [value1, value2, ...], ...]
                            Each inner array represents one row of the table, corresponding to the row headers in order.

                            Do NOT include any explanation, formatting, or code blocks — just return the raw JSON array.
                            """

    def format_prompt(self, user_prompt: str, row_headers: Optional[List[str]] = None,
                      column_headers: Optional[List[str]] = None, num_rows: int = 5, num_cols: int = 3) -> str:
        prompt = f"Generate table content for: {user_prompt}\n"
        if row_headers:
            prompt += f"Row headers: {', '.join(row_headers)}\n"
        if column_headers:
            prompt += f"Column headers: {', '.join(column_headers)}\n"
        if not row_headers and not column_headers:
            prompt += f"Generate a {num_rows}x{num_cols} table with appropriate values.\n"
        prompt += "\nReturn ONLY a JSON array of arrays."
        return prompt

    def extract_json_array(self, text: str) -> List[List[str]]:
        try:
            text = re.sub(r"```(?:json)?|```", "", text)
            start = text.find("[")
            end = text.rfind("]")
            if start == -1 or end == -1:
                return [[]]
            json_text = text[start:end + 1]
            data = json.loads(json_text)

            if isinstance(data, list) and all(isinstance(row, dict) for row in data):
                return [[str(cell) for cell in row.values()] for row in data]

            if isinstance(data[0], list) and isinstance(data[0][0], list):
                flat = []
                for chunk in data:
                    flat.extend(chunk)
                return flat

            return data
        except Exception as e:
            print("JSON parsing failed:", e)
            return [["N/A"] * 3 for _ in range(3)]

    def generate_table(self, prompt: str, row_headers: Optional[List[str]] = None,
                       column_headers: Optional[List[str]] = None,
                       num_rows: int = 5, num_cols: int = 3) -> List[List[str]]:
        full_prompt = self.format_prompt(prompt, row_headers, column_headers, num_rows, num_cols)
        input_text = f"<s>[INST] {self.system_prompt}\n{full_prompt} [/INST]"
        result = self.generator(input_text, max_new_tokens=1024, do_sample=True, temperature=0.7)[0]['generated_text']
        print("Raw LLM Output:\n", result)
        return self.extract_json_array(result)

    def save_to_docx(self, table_data: List[List[str]], file_path: str,
                     intro_text: Optional[str] = None, title: Optional[str] = None):
        doc = Document()
        if title:
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if intro_text:
            doc.add_paragraph(intro_text)
        rows = len(table_data)
        cols = len(table_data[0]) if rows > 0 else 0
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        for i, row in enumerate(table_data):
            for j, cell in enumerate(row):
                table.cell(i, j).text = str(cell)
        doc.save(file_path)
        print(f"Saved table to {file_path}")

        
        if file_path.endswith(".pdf"):
            docx_path = file_path.replace(".pdf", ".docx")
            self.save_to_docx(table_data, docx_path, intro_text, title)
            if DOCX2PDF_AVAILABLE:
                convert(docx_path, file_path)
                print(f"Converted to PDF: {file_path}")
            else:
                print("Cannot convert to PDF: docx2pdf not installed.")


def cli_main():
    parser = argparse.ArgumentParser(description="GenAI Table Filler")
    parser.add_argument("--prompt", type=str, help="Prompt for table content", required=False)
    parser.add_argument("--topic", type=str, help="High-level topic for automatic table generation", required=False)
    parser.add_argument("--columns", type=str, nargs="+", help="Column headers", required=False)
    parser.add_argument("--rows", type=str, nargs="+", help="Row headers", required=False)
    parser.add_argument("--output", type=str, choices=["console", "docx", "pdf"], default="console")
    parser.add_argument("--intro", type=str, default="AI-generated table based on your prompt.")
    parser.add_argument("--file", type=str, default="output.docx", help="Output file name (for docx/pdf)")
    args = parser.parse_args()

    filler = TableFiller()
    prompt = args.prompt or f"Generate a table about: {args.topic}" if args.topic else "Fill in the table."
    row_headers = args.rows
    col_headers = args.columns
    table = filler.generate_table(prompt, row_headers=row_headers, column_headers=col_headers)

    if args.output == "console":
        print("\nGenerated Table:\n")
        if col_headers:
            print(" | ".join([""] + col_headers))
        for i, row in enumerate(table):
            label = row_headers[i] if row_headers and i < len(row_headers) else f"Row {i+1}"
            print(f"{label} | " + " | ".join(row))
    else:
        file_ext = ".pdf" if args.output == "pdf" else ".docx"
        filename = args.file if args.file.endswith(file_ext) else args.file.replace(".docx", file_ext)
        table_with_headers = [[""] + col_headers] if col_headers else []
        if row_headers:
            table_with_headers += [[row_headers[i]] + row for i, row in enumerate(table)]
        else:
            table_with_headers += table
        filler.save_to_docx(table_with_headers, filename, intro_text=args.intro, title="Generated Table")


if __name__ == "__main__":
    cli_main()
